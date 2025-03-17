# -*- coding: utf-8 -*-
"""
Модуль для обработки документов и создания базы знаний
"""

import os
import re
import time
import hashlib
from datetime import datetime
from tqdm.auto import tqdm
import json

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

from cybersec_consultant.config import ConfigManager, DATA_DIR

class DocumentProcessor:
    """Класс для обработки документов различных форматов"""

    def __init__(self):
        """Инициализация обработчика документов"""
        self.supported_formats = {
            '.txt': self._process_txt,
            '.csv': self._process_csv,
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_docx,  # .doc будет обрабатываться как .docx
        }
        self.config_manager = ConfigManager()

    def _get_file_extension(self, filename):
        """Получает расширение файла в нижнем регистре"""
        _, ext = os.path.splitext(filename)
        return ext.lower()

    def _process_txt(self, file_bytes):
        """Обработка текстовых файлов"""
        return file_bytes.decode("utf-8")

    def _process_csv(self, file_bytes):
        """Обработка CSV файлов"""
        content = file_bytes.decode("utf-8")
        # Преобразуем CSV в более читаемый текстовый формат
        lines = content.split('\n')
        if not lines:
            return ""

        # Обрабатываем заголовок
        headers = lines[0].split(',')
        result = []

        # Обрабатываем данные
        for i, line in enumerate(lines[1:]):
            if not line.strip():
                continue
            values = line.split(',')
            row_text = f"Запись {i+1}:\n"
            for j, value in enumerate(values):
                if j < len(headers):
                    row_text += f"- {headers[j]}: {value}\n"
                else:
                    row_text += f"- Значение {j+1}: {value}\n"
            result.append(row_text)

        return "\n".join(result)

    def _process_pdf(self, file_bytes):
        """Обработка PDF файлов"""
        try:
            import io
            import PyPDF2

            pdf_file = io.BytesIO(file_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""

            # Проходим по всем страницам и извлекаем текст
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- Страница {page_num + 1} ---\n{page_text}\n"

            if not text.strip():
                return "PDF документ не содержит извлекаемого текста или является сканированным документом."
            return text
        except Exception as e:
            return f"Ошибка при обработке PDF файла: {str(e)}"

    def _process_docx(self, file_bytes):
        """Обработка DOCX файлов"""
        try:
            import io
            import docx

            docx_file = io.BytesIO(file_bytes)
            doc = docx.Document(docx_file)
            text = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)

            return "\n".join(text)
        except Exception as e:
            return f"Ошибка при обработке DOCX файла: {str(e)}"

    def process_file(self, file_path):
        """Обрабатывает файл и возвращает текстовое содержимое"""
        try:
            # Получаем расширение файла
            extension = self._get_file_extension(file_path)

            # Проверяем, поддерживается ли формат
            if extension not in self.supported_formats:
                supported = ", ".join(self.supported_formats.keys())
                raise ValueError(f"Формат файла {extension} не поддерживается. Поддерживаемые форматы: {supported}")

            # Читаем файл
            with open(file_path, 'rb') as f:
                file_bytes = f.read()

            # Обрабатываем файл в зависимости от формата
            print(f"Обработка файла формата {extension}...")
            text = self.supported_formats[extension](file_bytes)
            print(f"Файл обработан. Извлечено {len(text)} символов текста.")

            return text
        except Exception as e:
            print(f"Ошибка при обработке файла: {str(e)}")
            return None

class KnowledgeBaseManager:
    """Класс для управления базой знаний"""

    def __init__(self):
        """Инициализация менеджера базы знаний"""
        self.config_manager = ConfigManager()
        self.document_processor = DocumentProcessor()

        # Параметры из конфигурации
        self.chunk_size = self.config_manager.get_setting("settings", "chunk_size", 1024)
        self.chunk_overlap = self.config_manager.get_setting("settings", "chunk_overlap", 200)

        # Создаем директорию для базы знаний, если она не существует
        self.kb_dir = os.path.join(DATA_DIR, "knowledge_base")
        os.makedirs(self.kb_dir, exist_ok=True)

    def load_knowledge_base(self, file_path=None):
        """
        Загружает базу знаний по кибербезопасности из файла или использует демо-данные.

        Args:
            file_path (str): Путь к файлу с базой знаний

        Returns:
            str: Текст базы знаний
        """
        # Файл для базы знаний по умолчанию
        kb_file = os.path.join(self.kb_dir, "cybersecurity_kb.txt")

        # Если указан пользовательский файл, используем его
        if file_path:
            try:
                kb_text = self.document_processor.process_file(file_path)
                if kb_text:
                    # Сохраняем в файл по умолчанию для повторного использования
                    with open(kb_file, 'w', encoding='utf-8') as f:
                        f.write(kb_text)
                    print(f"✅ База знаний сохранена в файл: {kb_file}")
                    return kb_text
                else:
                    print("❌ Не удалось обработать файл с базой знаний.")
                    return self._generate_demo_knowledge_base()
            except Exception as e:
                print(f"❌ Ошибка при загрузке базы знаний: {str(e)}")
                return self._generate_demo_knowledge_base()

        # Если файл не указан, проверяем наличие сохраненной базы
        if os.path.exists(kb_file):
            print(f"📄 Найдена существующая база знаний: {kb_file}")

            # Спрашиваем пользователя, хочет ли он использовать существующий файл
            use_existing = input("Использовать существующую базу знаний? (y/n): ").lower().strip() == 'y'

            if use_existing:
                with open(kb_file, 'r', encoding='utf-8') as f:
                    kb_text = f.read()
                print(f"✅ Загружено {len(kb_text)} символов из существующего файла.")
                return kb_text

        # Если файл не существует или пользователь решил не использовать существующий
        print("⚠️ База знаний не найдена.")
        use_demo = input("Использовать демонстрационную базу знаний? (y/n): ").lower().strip() == 'y'

        if use_demo:
            return self._generate_demo_knowledge_base()
        else:
            raise ValueError("База знаний не предоставлена.")

    def _generate_demo_knowledge_base(self):
        """
        Создает демонстрационную базу знаний по кибербезопасности.

        Returns:
            str: Демонстрационный текст по кибербезопасности
        """
        demo_kb = """
# База знаний по кибербезопасности (демонстрационная версия)

## Аутентификация и авторизация

Аутентификация - это процесс подтверждения личности пользователя или системы. Авторизация - процесс определения доступа к ресурсам на основе подтвержденной личности.

### Методы аутентификации:

1. **Парольная аутентификация** - классический метод, основанный на "знании чего-то" (пароля). Наиболее распространен, но имеет множество уязвимостей.

2. **Двухфакторная аутентификация (2FA)** - повышает безопасность, требуя два различных фактора:
   - То, что вы знаете (пароль)
   - То, что у вас есть (телефон для получения SMS или приложение-аутентификатор)
   - То, чем вы являетесь (биометрия)

3. **Многофакторная аутентификация (MFA)** - использует три и более факторов для аутентификации.

4. **Биометрическая аутентификация** - использует уникальные биологические характеристики:
   - Отпечатки пальцев
   - Сканирование сетчатки глаза
   - Распознавание лица
   - Голосовая аутентификация

5. **Сертификаты и токены** - использование криптографических ключей и сертификатов.

### Лучшие практики аутентификации:

- Использование сложных паролей (не менее 12 символов, комбинация букв разного регистра, цифр и специальных символов)
- Регулярная смена паролей
- Внедрение двухфакторной аутентификации
- Блокировка аккаунта после определенного числа неудачных попыток
- Использование единого входа (SSO) для корпоративных систем
"""

        # Сохраняем демо-базу знаний для повторного использования
        kb_file = os.path.join(self.kb_dir, "cybersecurity_kb.txt")
        with open(kb_file, 'w', encoding='utf-8') as f:
            f.write(demo_kb)

        print(f"✅ Создана демонстрационная база знаний ({len(demo_kb)} символов).")
        return demo_kb

    def split_text_into_chunks(self, text):
        """
        Разбивает текст на чанки с перекрытием

        Args:
            text (str): Текст для разбиения

        Returns:
            list: Список документов (чанков)
        """
        print(f"🔄 Разбиваем текст на чанки (размер={self.chunk_size}, перекрытие={self.chunk_overlap})...")

        # Создаем разбиватель текста
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )

        # Разбиваем текст на чанки
        chunks = text_splitter.split_text(text)

        # Создаем документы с метаданными
        documents = []
        for i, chunk in enumerate(chunks):
            # Находим заголовки секций в чанке
            headers = re.findall(r'#+\s+(.+?)(?=\n|$)', chunk)
            categories = []

            # Определяем категории на основе заголовков
            for header in headers:
                header_lower = header.lower()
                if "аутентификац" in header_lower or "авторизац" in header_lower:
                    categories.append("аутентификация")
                elif "шифрован" in header_lower:
                    categories.append("шифрование")
                elif "сет" in header_lower or "экран" in header_lower or "ddos" in header_lower:
                    categories.append("сетевая_безопасность")
                elif "фишинг" in header_lower or "социальн" in header_lower or "угроз" in header_lower:
                    categories.append("угрозы")
                elif "вредонос" in header_lower or "защит" in header_lower:
                    categories.append("защита")

            # Если категории не найдены, используем "общее"
            if not categories:
                categories = ["общее"]

            # Создаем документ с метаданными
            doc = Document(
                page_content=chunk,
                metadata={
                    "chunk_id": i + 1,
                    "source": "База знаний по кибербезопасности",
                    "categories": categories,
                    "length": len(chunk)
                }
            )
            documents.append(doc)

        print(f"✅ Создано {len(documents)} чанков из текста ({len(text)} символов)")
        return documents

    def process_knowledge_base(self, file_path=None):
        """
        Выполняет полный процесс загрузки, обработки базы знаний

        Args:
            file_path (str): Путь к файлу с базой знаний (опционально)

        Returns:
            tuple: (loaded_text, documents)
        """
        print("=" * 80)
        print("🚀 ПРОЦЕСС ЗАГРУЗКИ И ПОДГОТОВКИ БАЗЫ ЗНАНИЙ ПО КИБЕРБЕЗОПАСНОСТИ")
        print("=" * 80)

        # 1. Загружаем базу знаний
        kb_text = self.load_knowledge_base(file_path)
        if not kb_text:
            print("❌ Не удалось загрузить базу знаний.")
            return None, None

        # 2. Разбиваем текст на чанки
        documents = self.split_text_into_chunks(kb_text)
        if not documents:
            print("❌ Не удалось разбить текст на чанки.")
            return kb_text, None

        print("\n" + "=" * 80)
        print("✅ БАЗА ЗНАНИЙ УСПЕШНО ЗАГРУЖЕНА И ПОДГОТОВЛЕНА")
        print(f"📚 Загружено: {len(kb_text)} символов текста")
        print(f"📑 Создано: {len(documents)} документов")
        print("=" * 80)

        return kb_text, documents
